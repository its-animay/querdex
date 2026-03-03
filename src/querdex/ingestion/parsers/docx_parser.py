from __future__ import annotations

from pathlib import Path

from docx import Document

from querdex.schemas import Section


class DOCXParser:
    source_format = "docx"

    def parse(self, path: Path, doc_id: str) -> list[Section]:
        doc = Document(str(path))
        sections: list[Section] = []
        current_heading = "Introduction"
        page_number = 1

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style is not None else ""
            if style_name.startswith("Heading"):
                current_heading = text
                continue

            sections.append(
                Section(
                    section_id=f"sec_{len(sections) + 1:04d}",
                    doc_id=doc_id,
                    content=text,
                    page_number=page_number,
                    source_format=self.source_format,
                    metadata={"heading": current_heading, "style": style_name},
                )
            )
            page_number += 1

        for table_index, table in enumerate(doc.tables, start=1):
            cells = []
            for row in table.rows:
                cells.append(" | ".join(cell.text.strip() for cell in row.cells))
            table_text = "\n".join(cells).strip()
            if table_text:
                sections.append(
                    Section(
                        section_id=f"sec_{len(sections) + 1:04d}",
                        doc_id=doc_id,
                        content=table_text,
                        page_number=page_number,
                        source_format=self.source_format,
                        metadata={
                            "heading": current_heading,
                            "type": "table",
                            "table_index": table_index,
                        },
                    )
                )
                page_number += 1

        return sections
