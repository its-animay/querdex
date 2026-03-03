from __future__ import annotations

import asyncio
from pathlib import Path

from docx import Document

from querdex.services import build_engine


def _create_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Revenue was 120 million dollars.")
    doc.add_heading("Liabilities", level=2)
    doc.add_paragraph("Liabilities were 50 million dollars.")
    doc.save(str(path))


def test_engine_index_and_query_docx(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    _create_docx(docx_path)
    db_path = tmp_path / "querdex.db"

    engine = build_engine(db_path)
    try:
        document_index = asyncio.run(engine.index_document(docx_path, doc_id="doc_docx"))
        assert document_index.doc_id == "doc_docx"

        result = engine.query_document(
            "doc_docx",
            "What are the liabilities?",
            session_id="docx_session",
        )

        assert result.source_nodes
        assert "No relevant content found" not in result.answer
    finally:
        engine.store.close()
