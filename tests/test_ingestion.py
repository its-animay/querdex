from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from querdex.ingestion import IngestionOrchestrator


def _write_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("This is the first paragraph.")
    doc.add_heading("Details", level=2)
    doc.add_paragraph("This is the second paragraph.")
    doc.save(str(path))


@pytest.mark.parametrize(
    ("filename", "content"),
    [
        ("sample.txt", "First block\n\nSecond block"),
        ("sample.md", "# Intro\n\nHello markdown"),
        ("sample.html", "<h1>Title</h1><p>Hello html</p>"),
    ],
)
def test_ingestion_conformance_text_like_formats(tmp_path: Path, filename: str, content: str) -> None:
    file_path = tmp_path / filename
    file_path.write_text(content, encoding="utf-8")

    orchestrator = IngestionOrchestrator()
    sections = orchestrator.parse(file_path, doc_id="doc_x")

    assert sections
    for section in sections:
        assert section.doc_id == "doc_x"
        assert section.page_number >= 1
        assert section.content.strip()
        assert section.source_format


def test_ingestion_conformance_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"
    _write_docx(file_path)

    orchestrator = IngestionOrchestrator()
    sections = orchestrator.parse(file_path, doc_id="doc_docx")

    assert sections
    assert all(section.doc_id == "doc_docx" for section in sections)
    assert all(section.source_format == "docx" for section in sections)


def test_ingestion_unsupported_extension(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.unsupported"
    file_path.write_text("data", encoding="utf-8")

    orchestrator = IngestionOrchestrator()
    with pytest.raises(ValueError):
        orchestrator.parse(file_path, doc_id="doc_bad")
